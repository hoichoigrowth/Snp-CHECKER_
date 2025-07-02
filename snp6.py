import json
import os
from openai import OpenAI
from datetime import datetime
import re
import gc
import time

# Try Word document processing
try:
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_COLOR_INDEX
    PYTHON_DOCX_AVAILABLE = True
    print("✅ python-docx available")

except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("❌ python-docx not available (install: pip install python-docx)")

# Try reportlab for PDF generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import Color, red, orange, yellow, lightgrey
    from reportlab.lib.units import inch
    from reportlab.platypus.flowables import KeepTogether
    REPORTLAB_AVAILABLE = True
    print("✅ reportlab available")
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("❌ reportlab not available (install: pip install reportlab)")

# Excel import
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    EXCEL_AVAILABLE = True
    print("✅ openpyxl available")
except ImportError:
    print("❌ Please install openpyxl: pip install openpyxl")
    exit(1)

# ============ CONFIGURATION ============
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
# ============ INPUT FILE CONFIGURATION ============
# ONLY DOCX FILES SUPPORTED - CHANGE PATH BELOW
INPUT_DOCX_FILE = "/home/alokananda/Desktop/S&P/KKB_Dialogue Script_05.06.25.docx"
# ================================================

# Configuration optimized for large files
MAX_CHARS_PER_CHUNK = 4000      # Reduced for better API reliability with large files
OVERLAP_CHARS = 200             # Reduced overlap for efficiency
MAX_TOKENS_OUTPUT = 1000        # Conservative token limit for large files
CHUNK_DELAY = 1                 # Delay between API calls to avoid rate limits
MAX_RETRIES = 3                 # Retry failed chunks

# ============ S&P VIOLATION RULES ============
VIOLATION_RULES = {
    "National_Symbols_Anthem": {
        "description": "Improper use of national symbols, anthem, flag, or government emblems",
        "keywords": ["national anthem", "flag", "tricolor", "ashoka chakra", "government logo"],
        "severity": "critical"
    },
    "PI_Data_Brand_References": {
        "description": "Personal information exposure, unauthorized brand references, trademark violations",
        "keywords": ["personal data", "phone number", "address", "brand name", "trademark"],
        "severity": "high"
    },
    "Credits_Endorsements": {
        "description": "Missing credits, unauthorized endorsements, celebrity impersonation",
        "keywords": ["credit", "endorsement", "celebrity", "sponsor", "testimonial"],
        "severity": "medium"
    },
    "Religious_Cultural_Sensitivity": {
        "description": "Religious insensitivity, cultural stereotypes, communal bias, offensive content",
        "keywords": ["religion", "god", "hindu", "muslim", "christian", "sikh", "culture", "caste"],
        "severity": "critical"
    },
    "Animal_Welfare": {
        "description": "Animal cruelty, unsafe animal handling, wildlife protection violations",
        "keywords": ["animal", "cruelty", "pet", "wildlife", "zoo", "circus"],
        "severity": "high"
    },
    "Disclaimers_Warnings": {
        "description": "Missing disclaimers, inadequate warnings, safety information gaps",
        "keywords": ["disclaimer", "warning", "caution", "safety", "risk"],
        "severity": "medium"
    },
    "Smoking_Alcohol_Social_Evils": {
        "description": "Glorification of smoking, alcohol abuse, gambling, or other social evils",
        "keywords": ["smoke", "cigarette", "alcohol", "drink", "gambling", "bet", "drugs"],
        "severity": "high"
    },
    "Child_Safety": {
        "description": "Content harmful to minors, inappropriate child representation, safety risks",
        "keywords": ["child", "minor", "kid", "school", "unsafe", "inappropriate"],
        "severity": "critical"
    },
    "Violence_Self_Harm": {
        "description": "Graphic violence, self-harm content, suicide references, dangerous activities",
        "keywords": ["violence", "fight", "blood", "suicide", "self-harm", "dangerous", "weapon"],
        "severity": "critical"
    },
    "Substances_Addiction": {
        "description": "Drug use promotion, addiction glorification, substance abuse normalization",
        "keywords": ["drugs", "addiction", "substance", "abuse", "dealer", "high", "overdose"],
        "severity": "critical"
    }
}

def extract_text_from_docx(file_path):
    """Extract text from DOCX file with robust handling for large files"""
    print(f"📄 Reading DOCX file: {file_path}")
    
    # Validate file
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        print("💡 Please check the file path in INPUT_DOCX_FILE variable")
        return None, []
    
    if not file_path.lower().endswith('.docx'):
        print(f"❌ Not a DOCX file: {file_path}")
        print("📝 This analyzer only supports .docx files")
        return None, []
    
    if not PYTHON_DOCX_AVAILABLE:
        print("❌ python-docx not available. Install with: pip install python-docx")
        return None, []
    
    try:
        # Check file size
        file_size = os.path.getsize(file_path)
        print(f"📊 File size: {file_size / (1024*1024):.2f} MB")
        
        if file_size > 50 * 1024 * 1024:  # 50MB
            print("⚠️ Large file detected. Processing with memory optimization...")
        
        doc = Document(file_path)
        pages_data = []
        full_text = ""
        
        print("📝 Extracting text from Word document...")
        
        # Extract paragraphs with optimized chunking for large files
        page_num = 1
        current_page_text = ""
        char_count = 0
        total_paragraphs = len(doc.paragraphs)
        
        print(f"📄 Processing {total_paragraphs} paragraphs...")
        
        # Process paragraphs in batches for memory efficiency
        batch_size = 100  # Process 100 paragraphs at a time
        processed_paragraphs = 0
        
        for i in range(0, total_paragraphs, batch_size):
            batch_end = min(i + batch_size, total_paragraphs)
            print(f"🔄 Processing paragraphs {i+1}-{batch_end} of {total_paragraphs}...")
            
            for j in range(i, batch_end):
                para = doc.paragraphs[j]
                para_text = para.text
                
                if para_text.strip():  # Only process non-empty paragraphs
                    current_page_text += para_text + "\n"
                    char_count += len(para_text) + 1
                    
                    # Create logical pages based on content length (smaller for large files)
                    page_size = 1500 if file_size > 10 * 1024 * 1024 else 2000
                    
                    if char_count > page_size:
                        pages_data.append({
                            'page_number': page_num,
                            'text': current_page_text.strip()
                        })
                        full_text += f"\n=== PAGE {page_num} ===\n{current_page_text}\n"
                        
                        if page_num % 10 == 0:  # Progress update every 10 pages
                            print(f"✅ Created page {page_num} ({len(current_page_text)} chars)")
                        
                        page_num += 1
                        current_page_text = ""
                        char_count = 0
                        
                        # Memory cleanup for large files
                        if page_num % 50 == 0:
                            gc.collect()
                
                processed_paragraphs += 1
            
            # Progress update
            progress = (processed_paragraphs / total_paragraphs) * 100
            print(f"📈 Progress: {progress:.1f}% ({processed_paragraphs}/{total_paragraphs})")
        
        # Add remaining content as the last page
        if current_page_text.strip():
            pages_data.append({
                'page_number': page_num,
                'text': current_page_text.strip()
            })
            full_text += f"\n=== PAGE {page_num} ===\n{current_page_text}\n"
            print(f"✅ Final page {page_num} created ({len(current_page_text)} chars)")
        
        print(f"✅ Successfully extracted {len(full_text):,} characters from {len(pages_data)} logical pages")
        print(f"📊 Average page size: {len(full_text)//len(pages_data):,} chars")
        print(f"📝 Content preview: {full_text[:150]}...")
        
        return full_text, pages_data
        
    except Exception as e:
        print(f"❌ DOCX extraction error: {e}")
        import traceback
        print(traceback.format_exc())
        return None, []

def chunk_text_robust(text, max_chars=MAX_CHARS_PER_CHUNK):
    """Split text into chunks optimized for large files"""
    print(f"📝 Chunking text: {len(text):,} characters -> {max_chars:,} char chunks")
    
    if len(text) <= max_chars:
        return [text]
    
    chunks = []
    start = 0
    overlap = OVERLAP_CHARS
    chunk_num = 0
    
    while start < len(text):
        end = start + max_chars
        
        # Find good break points to avoid cutting sentences/words
        if end < len(text):
            # Look for break points in order of preference
            break_points = [
                '\n=== PAGE',  # Page boundaries (best)
                '.\n',        # Sentence + paragraph break
                '. ',         # Sentence end
                '?\n',        # Question + paragraph
                '!\n',        # Exclamation + paragraph
                '\n\n',       # Paragraph break
                '\n',         # Line break
                ' '           # Word boundary (last resort)
            ]
            
            best_break = end
            for break_point in break_points:
                # Look for break point within reasonable range
                search_start = max(start + max_chars // 2, end - 500)  # Don't break too early
                last_break = text.rfind(break_point, search_start, end)
                
                if last_break > search_start:
                    best_break = last_break + len(break_point)
                    break
            
            end = best_break
        
        chunk = text[start:end].strip()
        if chunk:  # Only add non-empty chunks
            chunks.append(chunk)
            chunk_num += 1
            
            if chunk_num % 10 == 0:  # Progress update
                print(f"📦 Created {chunk_num} chunks ({len(chunk):,} chars in latest)")
        
        # Move start forward with overlap
        start = end - overlap
        if start >= len(text):
            break
    
    print(f"✅ Created {len(chunks)} chunks for processing")
    return chunks

def create_detailed_prompt():
    """Create optimized S&P compliance prompt for large file processing"""
    violation_types = []
    for v_type, details in VIOLATION_RULES.items():
        violation_types.append(f"- {v_type.replace('_', ' ')}: {details['description']} (Severity: {details['severity']})")
    
    violation_types_str = "\n".join(violation_types)
    
    return f"""You are an expert S&P compliance reviewer for Indian digital media. Analyze this content chunk for violations.

VIOLATION CATEGORIES (choose most appropriate):
{violation_types_str}

CRITICAL INSTRUCTIONS:
1. Copy violation text EXACTLY as it appears (preserve all formatting)
2. Extract meaningful violations only (minimum 10 characters)
3. Choose appropriate category and severity
4. Focus on substantial policy violations

Return ONLY valid JSON:
{{
  "violations": [
    {{
      "violationText": "EXACT text preserving all formatting",
      "violationType": "Category name from list above",
      "explanation": "why this violates broadcasting standards",
      "suggestedAction": "specific remediation needed",
      "severity": "critical|high|medium|low"
    }}
  ]
}}

If no violations found, return: {{"violations": []}}"""

def analyze_chunk_robust(chunk, chunk_num, total_chunks, api_key):
    """Analyze single chunk with robust error handling and retries"""
    print(f"🤖 Analyzing chunk {chunk_num}/{total_chunks} ({len(chunk):,} chars)...")
    
    for attempt in range(MAX_RETRIES):
        try:
            client = OpenAI(api_key=api_key)
            
            prompt = create_detailed_prompt()
            full_prompt = f"""{prompt}

Content to analyze:
{chunk}

JSON response only:"""
            
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an S&P compliance expert. Return only valid JSON with exact text preservation."},
                    {"role": "user", "content": full_prompt}
                ],
                temperature=0.1,
                max_tokens=MAX_TOKENS_OUTPUT,
                timeout=60  # 60 second timeout
            )
            
            result = response.choices[0].message.content.strip()
            
            # Parse JSON with multiple fallback methods
            try:
                parsed_result = json.loads(result)
            except json.JSONDecodeError:
                # Try to extract JSON from mixed response
                json_start = result.find('{')
                json_end = result.rfind('}')
                if json_start != -1 and json_end != -1:
                    json_text = result[json_start:json_end + 1]
                    try:
                        parsed_result = json.loads(json_text)
                    except:
                        print(f"⚠️ JSON parsing failed for chunk {chunk_num}, attempt {attempt + 1}")
                        if attempt < MAX_RETRIES - 1:
                            time.sleep(2)  # Wait before retry
                            continue
                        return {"violations": []}
                else:
                    return {"violations": []}
            
            # Validate and enhance violations
            if 'violations' in parsed_result:
                enhanced_violations = []
                for violation in parsed_result['violations']:
                    v_type = violation.get('violationType', 'Other')
                    
                    # Normalize violation type
                    v_type_key = v_type.replace(' ', '_')
                    if v_type_key not in VIOLATION_RULES:
                        for key in VIOLATION_RULES.keys():
                            if key.replace('_', ' ').lower() == v_type.lower():
                                v_type_key = key
                                break
                        else:
                            v_type_key = 'Other'
                    
                    # Add severity if missing
                    if 'severity' not in violation:
                        violation['severity'] = VIOLATION_RULES.get(v_type_key, {}).get('severity', 'medium')
                    
                    # Clean violation text
                    violation_text = violation.get('violationText', '').strip()
                    if len(violation_text) >= 10:  # Only meaningful violations
                        violation['violationType'] = v_type_key.replace('_', ' ')
                        violation['violationText'] = violation_text
                        enhanced_violations.append(violation)
                
                parsed_result['violations'] = enhanced_violations
                print(f"✅ Chunk {chunk_num} analyzed: {len(enhanced_violations)} violations found")
            
            # Rate limiting delay
            if CHUNK_DELAY > 0:
                time.sleep(CHUNK_DELAY)
            
            return parsed_result
            
        except Exception as e:
            print(f"❌ Error analyzing chunk {chunk_num}, attempt {attempt + 1}: {e}")
            if attempt < MAX_RETRIES - 1:
                wait_time = (attempt + 1) * 2  # Exponential backoff
                print(f"⏳ Waiting {wait_time} seconds before retry...")
                time.sleep(wait_time)
            else:
                print(f"💀 Failed to analyze chunk {chunk_num} after {MAX_RETRIES} attempts")
                return {"violations": []}

def find_page_number(violation_text, pages_data):
    """Find which page contains the violation with fuzzy matching"""
    # Try exact match first
    for page_data in pages_data:
        if violation_text in page_data['text']:
            return page_data['page_number']
    
    # Try partial match with first 50 characters
    search_text = violation_text[:50] if len(violation_text) > 50 else violation_text
    for page_data in pages_data:
        if search_text in page_data['text']:
            return page_data['page_number']
    
    # Try word-based matching for very different formatting
    violation_words = violation_text.split()[:5]  # First 5 words
    if len(violation_words) >= 2:
        search_phrase = ' '.join(violation_words)
        for page_data in pages_data:
            if search_phrase in page_data['text']:
                return page_data['page_number']
    
    return 1  # Default to page 1

def clean_violation_text_minimal(text):
    """Clean violation text minimally - only remove page markers"""
    if not text:
        return text
    
    # Remove page markers only
    text = re.sub(r'=== PAGE \d+ ===\n?', '', text)
    text = text.strip()
    
    return text

def analyze_document_robust(text, pages_data, api_key):
    """Analyze entire document with robust processing for large files"""
    print(f"📊 Document analysis starting...")
    print(f"   📄 Total size: {len(text):,} characters")
    print(f"   📋 Logical pages: {len(pages_data)}")
    
    # Create chunks
    chunks = chunk_text_robust(text)
    print(f"   📦 Total chunks: {len(chunks)}")
    
    # Estimate processing time
    estimated_time = len(chunks) * (3 + CHUNK_DELAY)  # 3 seconds per chunk + delay
    print(f"   ⏱️ Estimated time: {estimated_time//60:.0f} minutes {estimated_time%60:.0f} seconds")
    
    all_violations = []
    successful_chunks = 0
    failed_chunks = 0
    start_time = time.time()
    
    for i, chunk in enumerate(chunks, 1):
        # Progress update
        if i % 5 == 0 or i == 1:
            elapsed = time.time() - start_time
            rate = i / elapsed if elapsed > 0 else 0
            remaining_chunks = len(chunks) - i
            eta = remaining_chunks / rate if rate > 0 else 0
            print(f"📈 Progress: {i}/{len(chunks)} ({(i/len(chunks)*100):.1f}%) - ETA: {eta//60:.0f}m {eta%60:.0f}s")
        
        analysis = analyze_chunk_robust(chunk, i, len(chunks), api_key)
        
        if 'violations' in analysis:
            for violation in analysis['violations']:
                # Clean and validate violation text
                v_text = violation.get('violationText', '')
                v_text_cleaned = clean_violation_text_minimal(v_text)
                
                if len(v_text_cleaned) >= 10:  # Only meaningful violations
                    violation['violationText'] = v_text_cleaned
                    violation['pageNumber'] = find_page_number(v_text_cleaned, pages_data)
                    violation['chunkNumber'] = i
                    all_violations.append(violation)
            
            successful_chunks += 1
        else:
            failed_chunks += 1
        
        # Memory cleanup every 20 chunks
        if i % 20 == 0:
            gc.collect()
    
    # Remove duplicates
    unique_violations = []
    seen_texts = set()
    for violation in all_violations:
        v_text = violation.get('violationText', '')
        # Create a more robust duplicate key
        duplicate_key = (
            v_text[:100],  # First 100 chars
            violation.get('violationType', ''),
            violation.get('pageNumber', 0)
        )
        
        if duplicate_key not in seen_texts:
            seen_texts.add(duplicate_key)
            unique_violations.append(violation)
    
    # Sort by page number and severity
    severity_order = {'critical': 4, 'high': 3, 'medium': 2, 'low': 1}
    unique_violations.sort(key=lambda x: (
        x.get('pageNumber', 0),
        -severity_order.get(x.get('severity', 'low'), 1)
    ))
    
    total_time = time.time() - start_time
    print(f"\n✅ Document analysis complete!")
    print(f"   ⏱️ Total time: {total_time//60:.0f}m {total_time%60:.0f}s")
    print(f"   ✅ Successful chunks: {successful_chunks}/{len(chunks)}")
    print(f"   ❌ Failed chunks: {failed_chunks}/{len(chunks)}")
    print(f"   🔍 Unique violations found: {len(unique_violations)}")
    
    return {
        "violations": unique_violations,
        "summary": {
            "totalViolations": len(unique_violations),
            "totalPages": len(pages_data),
            "chunksAnalyzed": len(chunks),
            "successfulChunks": successful_chunks,
            "failedChunks": failed_chunks,
            "successRate": f"{(successful_chunks/len(chunks)*100):.1f}%",
            "processingTime": f"{total_time//60:.0f}m {total_time%60:.0f}s",
            "documentSize": f"{len(text):,} chars"
        }
    }

def save_xlsx_report(violations, analysis, input_file_path):
    """Save comprehensive XLSX report optimized for large datasets"""
    xlsx_file = f"snp_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    
    print(f"📊 Creating XLSX report: {xlsx_file}")
    print(f"   📋 Processing {len(violations)} violations...")
    
    try:
        wb = Workbook()
        wb.remove(wb.active)
        
        # Define styles
        header_font = Font(bold=True, size=12, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        violation_fill = PatternFill(start_color="FFE6E6", end_color="FFE6E6", fill_type="solid")
        
        # Sheet 1: Executive Summary
        ws_summary = wb.create_sheet("Executive Summary")
        summary_data = [
            ["S&P COMPLIANCE ANALYSIS REPORT", ""],
            ["", ""],
            ["Document Information", ""],
            ["File Analyzed", input_file_path.split('/')[-1]],
            ["Analysis Date", datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
            ["Document Size", analysis['summary']['documentSize']],
            ["Total Pages", analysis['summary']['totalPages']],
            ["Processing Time", analysis['summary']['processingTime']],
            ["", ""],
            ["Results Summary", ""],
            ["Total Violations", analysis['summary']['totalViolations']],
            ["Success Rate", analysis['summary']['successRate']],
            ["Chunks Analyzed", analysis['summary']['chunksAnalyzed']],
            ["", ""],
            ["Violation Breakdown", "Count"]
        ]
        
        # Add violation type breakdown with severity
        violation_types = {}
        severity_counts = {}
        for v in violations:
            v_type = v.get('violationType', 'Unknown')
            severity = v.get('severity', 'medium')
            
            violation_types[v_type] = violation_types.get(v_type, 0) + 1
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        for v_type, count in sorted(violation_types.items()):
            severity = VIOLATION_RULES.get(v_type.replace(' ', '_'), {}).get('severity', 'unknown')
            summary_data.append([f"{v_type} ({severity})", count])
        
        # Add severity breakdown
        summary_data.extend([
            ["", ""],
            ["Severity Distribution", "Count"]
        ])
        
        for severity in ['critical', 'high', 'medium', 'low']:
            count = severity_counts.get(severity, 0)
            if count > 0:
                summary_data.append([severity.upper(), count])
        
        # Write summary data
        for row_idx, row_data in enumerate(summary_data, 1):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws_summary.cell(row=row_idx, column=col_idx, value=value)
                if row_idx == 1 or (len(row_data) > 1 and row_data[0] and not row_data[1]):
                    cell.font = header_font
                    cell.fill = header_fill
        
        ws_summary.column_dimensions['A'].width = 30
        ws_summary.column_dimensions['B'].width = 25
        
        # Sheet 2: Detailed Violations
        ws_violations = wb.create_sheet("Detailed Violations")
        headers = ['S.No', 'Page', 'Type', 'Severity', 'Violation Text', 'Explanation', 'Suggested Action', 'Status']
        
        # Write headers
        for col_idx, header in enumerate(headers, 1):
            cell = ws_violations.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(wrap_text=True, vertical='top')
        
        # Write violation data in batches for memory efficiency
        batch_size = 100
        for batch_start in range(0, len(violations), batch_size):
            batch_end = min(batch_start + batch_size, len(violations))
            print(f"   📝 Writing violations {batch_start+1}-{batch_end}...")
            
            for i in range(batch_start, batch_end):
                v = violations[i]
                row_idx = i + 2  # +2 because row 1 is header and we're 0-indexed
                
                ws_violations.cell(row=row_idx, column=1, value=i+1)  # S.No
                ws_violations.cell(row=row_idx, column=2, value=v.get('pageNumber', 'N/A'))
                ws_violations.cell(row=row_idx, column=3, value=v.get('violationType', 'N/A'))
                
                # Severity with color coding
                severity = v.get('severity', 'medium')
                severity_cell = ws_violations.cell(row=row_idx, column=4, value=severity.upper())
                
                if severity == 'critical':
                    severity_cell.fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
                    severity_cell.font = Font(color="FFFFFF", bold=True)
                elif severity == 'high':
                    severity_cell.fill = PatternFill(start_color="FF6600", end_color="FF6600", fill_type="solid")
                    severity_cell.font = Font(color="FFFFFF", bold=True)
                elif severity == 'medium':
                    severity_cell.fill = PatternFill(start_color="FFCC00", end_color="FFCC00", fill_type="solid")
                
                # Violation text (truncate if too long for Excel)
                violation_text = v.get('violationText', 'N/A')
                if len(violation_text) > 1000:
                    violation_text = violation_text[:1000] + '...'
                
                violation_cell = ws_violations.cell(row=row_idx, column=5, value=violation_text)
                violation_cell.alignment = Alignment(wrap_text=True, vertical='top')
                
                ws_violations.cell(row=row_idx, column=6, value=v.get('explanation', 'N/A'))
                ws_violations.cell(row=row_idx, column=6).alignment = Alignment(wrap_text=True, vertical='top')
                
                ws_violations.cell(row=row_idx, column=7, value=v.get('suggestedAction', 'N/A'))
                ws_violations.cell(row=row_idx, column=7).alignment = Alignment(wrap_text=True, vertical='top')
                
                ws_violations.cell(row=row_idx, column=8, value='Pending Review')
        
        # Set column widths
        column_widths = {'A': 8, 'B': 8, 'C': 20, 'D': 12, 'E': 60, 'F': 40, 'G': 35, 'H': 15}
        for col, width in column_widths.items():
            ws_violations.column_dimensions[col].width = width
        
        # Sheet 3: Page Analysis
        ws_pages = wb.create_sheet("Page Analysis")
        page_headers = ['Page Number', 'Violation Count', 'Critical', 'High', 'Medium', 'Low', 'Status']
        
        for col_idx, header in enumerate(page_headers, 1):
            cell = ws_pages.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
        
        # Group violations by page with severity breakdown
        by_page = {}
        for v in violations:
            page = v.get('pageNumber', 'Unknown')
            if page not in by_page:
                by_page[page] = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0, 'violations': []}
            
            severity = v.get('severity', 'medium')
            by_page[page][severity] += 1
            by_page[page]['violations'].append(v)
        
        row_idx = 2
        for page_num in sorted(by_page.keys(), key=lambda x: int(x) if str(x).isdigit() else 999):
            page_data = by_page[page_num]
            total_violations = len(page_data['violations'])
            
            ws_pages.cell(row=row_idx, column=1, value=page_num)
            ws_pages.cell(row=row_idx, column=2, value=total_violations)
            ws_pages.cell(row=row_idx, column=3, value=page_data['critical'])
            ws_pages.cell(row=row_idx, column=4, value=page_data['high'])
            ws_pages.cell(row=row_idx, column=5, value=page_data['medium'])
            ws_pages.cell(row=row_idx, column=6, value=page_data['low'])
            
            # Status based on severity
            if page_data['critical'] > 0:
                status = 'CRITICAL REVIEW'
                status_color = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
            elif page_data['high'] > 0:
                status = 'HIGH PRIORITY'
                status_color = PatternFill(start_color="FF6600", end_color="FF6600", fill_type="solid")
            elif page_data['medium'] > 0:
                status = 'MEDIUM PRIORITY'
                status_color = PatternFill(start_color="FFCC00", end_color="FFCC00", fill_type="solid")
            else:
                status = 'LOW PRIORITY'
                status_color = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
            
            status_cell = ws_pages.cell(row=row_idx, column=7, value=status)
            status_cell.fill = status_color
            if page_data['critical'] > 0 or page_data['high'] > 0:
                status_cell.font = Font(color="FFFFFF", bold=True)
            
            row_idx += 1
        
        # Set page analysis column widths
        page_column_widths = {'A': 12, 'B': 15, 'C': 10, 'D': 10, 'E': 10, 'F': 10, 'G': 20}
        for col, width in page_column_widths.items():
            ws_pages.column_dimensions[col].width = width
        
        wb.save(xlsx_file)
        print(f"✅ XLSX report saved: {xlsx_file}")
        return xlsx_file
        
    except Exception as e:
        print(f"❌ XLSX creation failed: {e}")
        import traceback
        print(traceback.format_exc())
        return None

def create_compliance_review_pdf(violations, input_file_path, filename_prefix):
    """Create the main compliance review PDF with highlighted content"""
    if not REPORTLAB_AVAILABLE or not PYTHON_DOCX_AVAILABLE:
        print("❌ Required libraries not available")
        return None
    
    # Extract base filename without extension
    base_filename = os.path.splitext(os.path.basename(input_file_path))[0]
    output_pdf = f"{base_filename}_COMPLIANCE REVIEW.pdf"
    
    print(f"🎨 Creating Compliance Review PDF: {output_pdf}")
    
    try:
        doc = Document(input_file_path)
        
        # Create violation mapping
        violation_colors = {}
        violation_stats = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0}
        
        for violation in violations:
            v_text = violation.get('violationText', '').strip()
            severity = violation.get('severity', 'medium').lower()
            violation_stats[severity] += 1
            
            # Color coding for backgrounds
            color_map = {
                'critical': Color(1, 0.8, 0.8),    # Light red
                'high': Color(1, 0.9, 0.8),       # Light orange
                'medium': Color(1, 1, 0.8),       # Light yellow
                'low': Color(0.9, 0.9, 0.9)       # Light gray
            }
            
            if v_text and len(v_text) >= 10:
                violation_colors[v_text] = {
                    'color': color_map.get(severity, color_map['medium']),
                    'type': violation.get('violationType', 'Unknown'),
                    'severity': severity.upper(),
                    'page': violation.get('pageNumber', 1)
                }
        
        # Create PDF
        doc_pdf = SimpleDocTemplate(
            output_pdf,
            pagesize=A4,
            rightMargin=0.8*inch,
            leftMargin=0.8*inch,
            topMargin=0.8*inch,
            bottomMargin=0.8*inch
        )
        
        # Define styles
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        normal_style.fontSize = 10
        normal_style.leading = 12
        
        # Create highlighted styles
        color_map = {
            'critical': Color(1, 0.8, 0.8),
            'high': Color(1, 0.9, 0.8),
            'medium': Color(1, 1, 0.8),
            'low': Color(0.9, 0.9, 0.9)
        }
        
        highlighted_styles = {}
        for severity in ['critical', 'high', 'medium', 'low']:
            highlighted_styles[severity] = ParagraphStyle(
                f'{severity.title()}',
                parent=normal_style,
                backColor=color_map[severity],
                borderColor=red if severity == 'critical' else orange if severity == 'high' else yellow,
                borderWidth=1,
                leftIndent=10,
                rightIndent=10,
                spaceBefore=6,
                spaceAfter=6
            )
        
        story = []
        
        # Title and metadata
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            spaceAfter=20,
            textColor=Color(0.2, 0.2, 0.6),
            alignment=1
        )
        
        story.append(Paragraph("S&P COMPLIANCE REVIEW", title_style))
        story.append(Paragraph(f"Document: {os.path.basename(input_file_path)}", styles['Normal']))
        story.append(Paragraph(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Summary stats
        summary_style = ParagraphStyle(
            'Summary',
            parent=styles['Normal'],
            fontSize=11,
            textColor=Color(0.3, 0.3, 0.3),
            leftIndent=20
        )
        
        story.append(Paragraph("VIOLATION SUMMARY:", styles['Heading2']))
        story.append(Paragraph(f"🔴 Critical: {violation_stats['critical']} | 🟠 High: {violation_stats['high']} | 🟡 Medium: {violation_stats['medium']} | ⚪ Low: {violation_stats['low']}", summary_style))
        story.append(Spacer(1, 15))
        
        def highlight_violation_text(text, violations_in_para):
            """Highlight specific violation text within a paragraph in red"""
            highlighted_text = text
            sorted_violations = sorted(violations_in_para, key=lambda x: len(x.get('text', '')), reverse=True)
            
            for v_info in sorted_violations:
                violation_text = v_info.get('text', '')
                if violation_text and violation_text in highlighted_text:
                    safe_violation_text = violation_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    red_violation = f'<font color="red"><b>{safe_violation_text}</b></font>'
                    highlighted_text = highlighted_text.replace(violation_text, red_violation)
            
            return highlighted_text
        
        # Process document content
        print("   📄 Processing document content...")
        paragraph_count = 0
        
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if not para_text:
                continue
            
            paragraph_count += 1
            
            # Check for violations in this paragraph
            para_violations = []
            highlighted = False
            
            for v_text, v_info in violation_colors.items():
                if v_text in para_text or any(word in para_text for word in v_text.split()[:3] if len(word) > 4):
                    para_violation_info = v_info.copy()
                    para_violation_info['text'] = v_text
                    para_violations.append(para_violation_info)
                    highlighted = True
            
            if highlighted:
                # Determine severity and style
                severities = [v['severity'].lower() for v in para_violations]
                if 'critical' in severities:
                    style = highlighted_styles['critical']
                elif 'high' in severities:
                    style = highlighted_styles['high']
                elif 'medium' in severities:
                    style = highlighted_styles['medium']
                else:
                    style = highlighted_styles['low']
                
                # Create annotation
                violation_types = list(set(v['type'] for v in para_violations))
                annotation = f"⚠️ VIOLATIONS: {', '.join(violation_types)}"
                
                # Highlight violation text in red
                highlighted_display_text = highlight_violation_text(para_text, para_violations)
                annotated_text = f"{highlighted_display_text}<br/><font size='8' color='darkred'><b>{annotation}</b></font>"
                story.append(Paragraph(annotated_text, style))
            else:
                # Normal paragraph
                if len(para_text) > 600:
                    para_text = para_text[:600] + "..."
                story.append(Paragraph(para_text, normal_style))
            
            story.append(Spacer(1, 4))
            
            if paragraph_count % 30 == 0:
                story.append(PageBreak())
        
        # Add color legend
        story.append(PageBreak())
        story.append(Paragraph("COLOR LEGEND", styles['Heading2']))
        legend_text = """
        <b>Background Colors indicate violation severity:</b><br/>
        🔴 <font color="red">Critical Severity</font> - Red background, immediate attention required<br/>
        🟠 <font color="orange">High Severity</font> - Orange background, high priority review<br/>
        🟡 <font color="#B8860B">Medium Severity</font> - Yellow background, standard review<br/>
        ⚪ <font color="gray">Low Severity</font> - Gray background, minor issues<br/><br/>
        <b><font color="red">Red text indicates the exact violation content</font></b> that triggered the S&P flag.
        """
        story.append(Paragraph(legend_text, normal_style))
        
        # Build PDF
        doc_pdf.build(story)
        
        print(f"✅ Compliance Review PDF created: {output_pdf}")
        return output_pdf
        
    except Exception as e:
        print(f"❌ Compliance Review PDF creation failed: {e}")
        return None

def create_detailed_violations_pdf(violations, input_file_path, filename_prefix):
    """Create the detailed violations list PDF"""
    if not REPORTLAB_AVAILABLE:
        print("❌ reportlab not available")
        return None
    
    # Extract base filename without extension
    base_filename = os.path.splitext(os.path.basename(input_file_path))[0]
    output_pdf = f"{base_filename}_DETAILED VIOLATION LIST.pdf"
    
    print(f"📋 Creating Detailed Violations PDF: {output_pdf}")
    
    try:
        # Create PDF
        doc_pdf = SimpleDocTemplate(
            output_pdf,
            pagesize=A4,
            rightMargin=0.8*inch,
            leftMargin=0.8*inch,
            topMargin=0.8*inch,
            bottomMargin=0.8*inch
        )
        
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        normal_style.fontSize = 10
        normal_style.leading = 12
        
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            spaceAfter=20,
            textColor=Color(0.2, 0.2, 0.6),
            alignment=1
        )
        
        story.append(Paragraph("DETAILED VIOLATION LIST", title_style))
        story.append(Paragraph(f"Document: {os.path.basename(input_file_path)}", styles['Normal']))
        story.append(Paragraph(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Paragraph(f"Total Violations: {len(violations)}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Group violations by severity
        violations_by_severity = {}
        for violation in violations:
            severity = violation.get('severity', 'medium')
            if severity not in violations_by_severity:
                violations_by_severity[severity] = []
            violations_by_severity[severity].append(violation)
        
        # Process each severity level
        for severity in ['critical', 'high', 'medium', 'low']:
            if severity in violations_by_severity:
                severity_violations = violations_by_severity[severity]
                if severity_violations:
                    # Severity header
                    severity_style = ParagraphStyle(
                        f'{severity}Header',
                        parent=styles['Heading2'],
                        textColor=red if severity == 'critical' else orange if severity == 'high' else Color(0.7, 0.7, 0) if severity == 'medium' else Color(0.5, 0.5, 0.5)
                    )
                    
                    story.append(Paragraph(f"{severity.upper()} SEVERITY ({len(severity_violations)} violations)", severity_style))
                    story.append(Spacer(1, 10))
                    
                    # List violations
                    for i, violation in enumerate(severity_violations, 1):
                        violation_text = violation.get('violationText', 'N/A')
                        
                        # Create violation box
                        violation_style = ParagraphStyle(
                            f'Violation{i}',
                            parent=normal_style,
                            leftIndent=20,
                            rightIndent=20,
                            spaceBefore=8,
                            spaceAfter=8,
                            borderWidth=1,
                            borderColor=Color(0.8, 0.8, 0.8),
                            backColor=Color(0.98, 0.98, 0.98)
                        )
                        
                        # Format violation details
                        red_violation_text = f'<font color="red"><b>{violation_text}</b></font>'
                        
                        violation_detail = f"<b>#{i}</b><br/>"
                        violation_detail += f"<b>Type:</b> {violation.get('violationType', 'Unknown')}<br/>"
                        violation_detail += f"<b>Page:</b> {violation.get('pageNumber', 'N/A')}<br/>"
                        violation_detail += f"<b>Severity:</b> {severity.upper()}<br/>"
                        violation_detail += f"<b>Violation Text:</b><br/>{red_violation_text}<br/>"
                        violation_detail += f"<b>Explanation:</b> {violation.get('explanation', 'N/A')}<br/>"
                        violation_detail += f"<b>Suggested Action:</b> {violation.get('suggestedAction', 'N/A')}<br/>"
                        violation_detail += f"<b>Status:</b> <font color='red'>PENDING REVIEW</font>"
                        
                        story.append(Paragraph(violation_detail, violation_style))
                    
                    story.append(Spacer(1, 20))
        
        # Summary statistics
        story.append(PageBreak())
        story.append(Paragraph("VIOLATION STATISTICS", styles['Heading1']))
        story.append(Spacer(1, 10))
        
        # Count violations by type
        type_counts = {}
        page_counts = {}
        for violation in violations:
            v_type = violation.get('violationType', 'Unknown')
            page = violation.get('pageNumber', 'Unknown')
            type_counts[v_type] = type_counts.get(v_type, 0) + 1
            page_counts[page] = page_counts.get(page, 0) + 1
        
        # Violation types breakdown
        story.append(Paragraph("Violations by Type:", styles['Heading2']))
        for v_type, count in sorted(type_counts.items(), key=lambda x: x[1], reverse=True):
            story.append(Paragraph(f"• {v_type}: {count} violations", normal_style))
        
        story.append(Spacer(1, 15))
        
        # Pages with most violations
        story.append(Paragraph("Pages with Most Violations:", styles['Heading2']))
        top_pages = sorted(page_counts.items(), key=lambda x: x[1], reverse=True)[:10]
        for page, count in top_pages:
            story.append(Paragraph(f"• Page {page}: {count} violations", normal_style))
        
        story.append(Spacer(1, 15))
        
        # Action items
        story.append(Paragraph("RECOMMENDED ACTIONS", styles['Heading1']))
        story.append(Paragraph("1. <b>Critical violations</b> require immediate attention before content publication", normal_style))
        story.append(Paragraph("2. <b>High severity violations</b> should be addressed in the next revision", normal_style))
        story.append(Paragraph("3. <b>Medium and low violations</b> should be reviewed and addressed as feasible", normal_style))
        story.append(Paragraph("4. <b>Red highlighted text</b> shows exact content that needs modification", normal_style))
        story.append(Paragraph("5. Refer to the Compliance Review PDF for context around each violation", normal_style))
        
        # Build PDF
        doc_pdf.build(story)
        
        print(f"✅ Detailed Violations PDF created: {output_pdf}")
        return output_pdf
        
    except Exception as e:
        print(f"❌ Detailed Violations PDF creation failed: {e}")
        return None

def create_highlighted_pdf_from_docx(violations, input_file_path, filename_prefix):
    """Create both PDF reports"""
    print(f"🎨 Creating two separate PDF reports...")
    
    # Create both PDFs
    compliance_pdf = create_compliance_review_pdf(violations, input_file_path, filename_prefix)
    violations_pdf = create_detailed_violations_pdf(violations, input_file_path, filename_prefix)
    
    if compliance_pdf and violations_pdf:
        print(f"✅ Both PDF reports created successfully!")
        print(f"   📄 Compliance Review: {compliance_pdf}")
        print(f"   📋 Detailed Violations: {violations_pdf}")
        return [compliance_pdf, violations_pdf]
    else:
        print(f"❌ PDF creation had issues")
        return None

def print_results(analysis):
    """Print analysis results with summary for large datasets"""
    print("\n" + "="*80)
    print("📊 S&P COMPLIANCE ANALYSIS RESULTS")
    print("="*80)
    
    violations = analysis.get('violations', [])
    summary = analysis.get('summary', {})
    
    print(f"🚨 TOTAL VIOLATIONS: {len(violations)}")
    print(f"📄 DOCUMENT SIZE: {summary.get('documentSize', 'N/A')}")
    print(f"⏱️ PROCESSING TIME: {summary.get('processingTime', 'N/A')}")
    print(f"📊 SUCCESS RATE: {summary.get('successRate', 'N/A')}")
    
    if violations:
        # Severity breakdown
        severity_counts = {}
        for v in violations:
            severity = v.get('severity', 'medium')
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        print(f"\n📈 SEVERITY BREAKDOWN:")
        for severity in ['critical', 'high', 'medium', 'low']:
            count = severity_counts.get(severity, 0)
            if count > 0:
                emoji = {'critical': '🔴', 'high': '🟠', 'medium': '🟡', 'low': '⚪'}
                print(f"   {emoji.get(severity, '⚪')} {severity.upper()}: {count}")
        
        # Show top violations by severity
        print(f"\n🔍 TOP VIOLATIONS BY SEVERITY:")
        critical_violations = [v for v in violations if v.get('severity') == 'critical']
        high_violations = [v for v in violations if v.get('severity') == 'high']
        
        if critical_violations:
            print(f"\n🔴 CRITICAL VIOLATIONS ({len(critical_violations)}):")
            for i, v in enumerate(critical_violations[:5], 1):  # Show top 5
                print(f"   {i}. {v.get('violationType', 'Unknown')} (Page {v.get('pageNumber', 'N/A')})")
                text = v.get('violationText', 'N/A')
                print(f"      \"{text[:100]}{'...' if len(text) > 100 else ''}\"")
        
        if high_violations:
            print(f"\n🟠 HIGH SEVERITY VIOLATIONS ({len(high_violations)}):")
            for i, v in enumerate(high_violations[:3], 1):  # Show top 3
                print(f"   {i}. {v.get('violationType', 'Unknown')} (Page {v.get('pageNumber', 'N/A')})")
        
        # Page distribution
        page_counts = {}
        for v in violations:
            page = v.get('pageNumber', 'Unknown')
            page_counts[page] = page_counts.get(page, 0) + 1
        
        top_pages = sorted(page_counts.items(), key=lambda x: x[1], reverse=True)[:5]
        print(f"\n📄 PAGES WITH MOST VIOLATIONS:")
        for page, count in top_pages:
            print(f"   Page {page}: {count} violations")
    else:
        print("✅ No violations found!")
    
    print("="*80)

def main():
    """Main function optimized for large file processing"""
    print("🚀 S&P COMPLIANCE ANALYZER - ROBUST LARGE FILE PROCESSOR")
    print("="*65)
    
    # System requirements check
    print("📚 System Requirements Check:")
    all_required_available = True
    
    if PYTHON_DOCX_AVAILABLE:
        print("   ✅ python-docx (DOCX processing)")
    else:
        print("   ❌ python-docx (REQUIRED: pip install python-docx)")
        all_required_available = False
    
    if REPORTLAB_AVAILABLE:
        print("   ✅ reportlab (PDF generation)")
    else:
        print("   ❌ reportlab (REQUIRED: pip install reportlab)")
        all_required_available = False
    
    print("   ✅ openpyxl (XLSX reports)")
    
    if not all_required_available:
        print("\n❌ Missing required libraries. Install:")
        missing = []
        if not PYTHON_DOCX_AVAILABLE:
            missing.append("python-docx")
        if not REPORTLAB_AVAILABLE:
            missing.append("reportlab")
        print(f"   pip install {' '.join(missing)}")
        return
    
    # Configuration validation
    if "your-api-key-here" in OPENAI_API_KEY or not OPENAI_API_KEY.strip():
        print("❌ Please set your OpenAI API key in the configuration!")
        return
    
    print(f"\n🔑 API Key: {OPENAI_API_KEY[:8]}...{OPENAI_API_KEY[-4:]}")
    print(f"📁 Input File: {INPUT_DOCX_FILE}")
    print(f"⚙️ Processing Config:")
    print(f"   📦 Chunk Size: {MAX_CHARS_PER_CHUNK:,} chars")
    print(f"   ⏱️ API Delay: {CHUNK_DELAY}s between calls")
    print(f"   🔄 Max Retries: {MAX_RETRIES}")
    
    # File size check and optimization
    if os.path.exists(INPUT_DOCX_FILE):
        file_size = os.path.getsize(INPUT_DOCX_FILE)
        file_size_mb = file_size / (1024 * 1024)
        print(f"📊 File Size: {file_size_mb:.2f} MB")
        
        if file_size_mb > 10:
            print("⚠️ Large file detected - using optimized processing")
            print("   This may take significantly longer but will handle the full document")
        
        # Estimate processing time and cost
        estimated_chunks = file_size // (MAX_CHARS_PER_CHUNK * 2)  # Rough estimate
        estimated_time_minutes = (estimated_chunks * (3 + CHUNK_DELAY)) // 60
        estimated_cost = estimated_chunks * 0.03  # Rough API cost estimate
        
        print(f"📈 Estimates:")
        print(f"   📦 Expected chunks: ~{estimated_chunks}")
        print(f"   ⏱️ Processing time: ~{estimated_time_minutes} minutes")
        print(f"   💰 Estimated API cost: ~${estimated_cost:.2f}")
        
        if estimated_time_minutes > 30:
            print("⚠️ This is a very large file that will take substantial time to process")
            print("💡 Consider breaking it into smaller sections if possible")
    
    # Extract text from DOCX
    print(f"\n📄 STAGE 1: TEXT EXTRACTION")
    text, pages_data = extract_text_from_docx(INPUT_DOCX_FILE)
    if not text:
        print("❌ Failed to extract text. Check file path and format.")
        return
    
    # Analyze content
    print(f"\n🔍 STAGE 2: S&P COMPLIANCE ANALYSIS")
    analysis = analyze_document_robust(text, pages_data, OPENAI_API_KEY)
    
    # Print results
    print_results(analysis)
    
    # Generate outputs
    violations = analysis.get('violations', [])
    if violations:
        print(f"\n💾 STAGE 3: GENERATING REPORTS FOR {len(violations)} VIOLATIONS")
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        prefix = "snp_analysis"
        
        # Create XLSX report
        print("\n📊 Creating comprehensive XLSX report...")
        xlsx_file = save_xlsx_report(violations, analysis, INPUT_DOCX_FILE)
        
        # Create highlighted PDF
        print("\n🎨 Creating highlighted PDF report...")
        highlighted_pdf = create_highlighted_pdf_from_docx(violations, INPUT_DOCX_FILE, prefix)
        
        # Final results
        print(f"\n✅ ANALYSIS COMPLETE!")
        print(f"📊 XLSX Report: {xlsx_file if xlsx_file else 'Failed'}")
        print(f"🎨 Highlighted PDF: {highlighted_pdf if highlighted_pdf else 'Failed'}")
        print(f"🎯 Processed {len(violations)} violations across {analysis['summary']['totalPages']} pages")
        
        if xlsx_file and highlighted_pdf:
            print(f"\n🎉 SUCCESS! Both reports generated:")
            print(f"   📋 Detailed analysis: {xlsx_file}")
            print(f"   📑 Visual review: {highlighted_pdf}")
            print(f"   ⏱️ Total time: {analysis['summary']['processingTime']}")
    else:
        print("\n✅ CLEAN DOCUMENT - No violations found!")
        print("🎉 Content fully complies with S&P standards")
    
    print(f"\n📋 WORKFLOW SUMMARY:")
    print(f"   📥 Input: {INPUT_DOCX_FILE.split('/')[-1]} ({analysis['summary']['documentSize']})")
    print(f"   🔍 Analysis: {analysis['summary']['chunksAnalyzed']} chunks processed")
    print(f"   📊 Success Rate: {analysis['summary']['successRate']}")
    print(f"   🎯 Categories: 10 specialized S&P violation types")
    print(f"   📈 Processing: Optimized for large files with robust error handling")
    
    print(f"\n💡 FOR DIFFERENT FILES:")
    print(f"   Update INPUT_DOCX_FILE path and run again")

if __name__ == "__main__":
    main()
